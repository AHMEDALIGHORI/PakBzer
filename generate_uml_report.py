#!/usr/bin/env python3
"""Generate UML diagrams and embed them in PakBzer OOP report (.doc + .docx)."""
import base64
import re
from pathlib import Path

import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyArrowPatch, Polygon, FancyBboxPatch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).parent
DIAGRAMS = ROOT / "diagrams"
DOC_HTML = ROOT / "PakBzer_OOP_Report.doc"
DOCX_OUT = ROOT / "PakBzer_OOP_Report.docx"


def draw_class_box(ax, x, y, w, h, name, lines, abstract=False, iface=False):
    """Draw a UML class/interface box."""
    style = dict(boxstyle="round,pad=0.02", linewidth=1.2, edgecolor="#01411C", facecolor="#f1f6f2")
    if abstract:
        name = f"<<abstract>> {name}"
    if iface:
        style["facecolor"] = "#e8f4ec"
        name = f"<<interface>> {name}"
    box = FancyBboxPatch((x, y), w, h, **style)
    ax.add_patch(box)
    ax.plot([x, x + w], [y + h - 0.35, y + h - 0.35], color="#01411C", linewidth=1)
    ax.text(x + w / 2, y + h - 0.18, name, ha="center", va="center", fontsize=8, fontweight="bold")
    body = "\n".join(lines)
    ax.text(x + 0.08, y + h - 0.55, body, ha="left", va="top", fontsize=6.5, family="monospace")


def draw_inheritance(ax, child_cx, child_top_y, parent_cx, parent_bottom_y):
    """Hollow triangle at parent — child extends parent."""
    tip_y = parent_bottom_y
    base_y = tip_y - 0.28
    tri = Polygon(
        [(parent_cx, tip_y), (parent_cx - 0.18, base_y), (parent_cx + 0.18, base_y)],
        closed=True, fill=False, edgecolor="#046A38", linewidth=1.5,
    )
    ax.add_patch(tri)
    ax.plot([child_cx, parent_cx], [child_top_y, base_y], color="#046A38", linewidth=1.5)


def draw_composition(ax, owner_x, owner_y, part_x, part_y, label="1..*"):
    """Filled diamond on owner side (composition)."""
    dx, dy = part_x - owner_x, part_y - owner_y
    length = (dx ** 2 + dy ** 2) ** 0.5
    ux, uy = dx / length, dy / length
    diamond_cx = owner_x + ux * 0.22
    diamond_cy = owner_y + uy * 0.22
    px, py = -uy * 0.1, ux * 0.1
    diamond = Polygon(
        [
            (diamond_cx + ux * 0.14, diamond_cy + uy * 0.14),
            (diamond_cx + px, diamond_cy + py),
            (diamond_cx - ux * 0.14, diamond_cy - uy * 0.14),
            (diamond_cx - px, diamond_cy - py),
        ],
        closed=True, facecolor="#01411C", edgecolor="#01411C",
    )
    ax.add_patch(diamond)
    ax.plot([owner_x + ux * 0.28, part_x], [owner_y + uy * 0.28, part_y], color="#01411C", linewidth=1.5)
    mx, my = (owner_x + part_x) / 2, (owner_y + part_y) / 2
    ax.text(mx, my + 0.12, label, fontsize=6.5, ha="center", color="#333")


def draw_association(ax, x1, y1, x2, y2, label="", arrow_end=False):
    style = "-|>" if arrow_end else "-"
    ax.annotate(
        "", xy=(x2, y2), xytext=(x1, y1),
        arrowprops=dict(arrowstyle=style, color="#333", lw=1.3, shrinkA=2, shrinkB=2),
    )
    if label:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.15, label, fontsize=6.5, ha="center")


def draw_dependency(ax, x1, y1, x2, y2, label=""):
    ax.annotate(
        "", xy=(x2, y2), xytext=(x1, y1),
        arrowprops=dict(arrowstyle="-|>", color="#666", lw=1.2, linestyle="dashed", shrinkA=2, shrinkB=2),
    )
    if label:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.15, label, fontsize=6.5, ha="center", color="#666")


def draw_realization(ax, impl_x, impl_y, iface_x, iface_y):
    """Dashed line + hollow triangle at interface (implements)."""
    ax.plot([impl_x, iface_x], [impl_y, iface_y + 0.35], color="#046A38", linewidth=1.2, linestyle="dashed")
    tri = Polygon(
        [(iface_x, iface_y), (iface_x - 0.16, iface_y + 0.26), (iface_x + 0.16, iface_y + 0.26)],
        closed=True, fill=False, edgecolor="#046A38", linewidth=1.2, linestyle="dashed",
    )
    ax.add_patch(tri)


def create_class_diagram(path: Path):
    fig, ax = plt.subplots(figsize=(14, 10))
    ax.set_xlim(0, 14)
    ax.set_ylim(-1.4, 10)
    ax.axis("off")
    ax.set_title("PakBzer — UML Class Diagram (Domain + Service Layer)", fontsize=13, fontweight="bold", color="#01411C", pad=12)

    # BaseEntity
    draw_class_box(ax, 5.5, 8.2, 3.0, 1.2, "BaseEntity", ["- id: Long", "- createdAt, updatedAt", "+ getId()"], abstract=True)

    # Subclasses
    draw_class_box(ax, 1.0, 5.8, 2.8, 1.6, "Product", ["- name, price, stock", "- category: Category", "+ getAverageRating()", "+ isInStock()"])
    draw_class_box(ax, 5.5, 5.8, 2.6, 1.4, "User", ["- email, fullName", "- password, role: Role", "- provider"])
    draw_class_box(ax, 10.0, 5.8, 2.8, 1.4, "Review", ["- authorName", "- rating, comment", "- product: Product"])

    draw_inheritance(ax, 2.4, 5.8, 6.2, 8.2)
    draw_inheritance(ax, 6.8, 5.8, 7.0, 8.2)
    draw_inheritance(ax, 11.4, 5.8, 7.8, 8.2)

    # Enums
    draw_class_box(ax, 0.2, 3.5, 2.2, 1.0, "Category", ["MEN, WOMEN, KIDS", "BOY, ADULT"], iface=False)
    draw_class_box(ax, 5.8, 3.5, 2.0, 0.8, "Role", ["USER, ADMIN"])
    draw_association(ax, 2.4, 5.8, 1.3, 4.5, "category")
    draw_association(ax, 6.8, 5.8, 6.8, 4.3, "role")

    # Product-Review composition
    draw_composition(ax, 3.8, 6.5, 10.0, 6.5, "1..*")

    # Cart composition
    draw_class_box(ax, 8.5, 2.0, 2.4, 1.2, "Cart", ["- items: List", "+ add(Product)", "+ getTotalPrice()"])
    draw_class_box(ax, 11.5, 2.0, 2.2, 1.2, "CartItem", ["- productId, qty", "+ getLineTotal()"])
    draw_composition(ax, 10.9, 2.6, 11.5, 2.6, "1..*")

    # Service layer
    draw_class_box(ax, 0.5, 0.3, 2.8, 1.0, "ProductService", ["+ findAll()", "+ findById()", "+ search()"], iface=True)
    draw_class_box(ax, 3.8, 0.3, 3.0, 1.0, "ProductServiceImpl", ["- productRepository", "+ search()"])
    draw_realization(ax, 5.3, 1.3, 2.0, 1.3)

    draw_class_box(ax, 7.5, 0.3, 2.8, 1.0, "ProductRepository", ["+ findByCategory()", "+ findAll()"], iface=True)
    draw_dependency(ax, 5.3, 0.3, 8.9, 1.3, "<<uses>>")

    draw_class_box(ax, 10.8, 0.3, 2.8, 1.0, "CustomOidcUserService", ["+ loadUser()"])
    draw_class_box(ax, 10.8, -1.0, 2.8, 0.7, "OidcUserService", ["+ loadUser()"], abstract=True)
    draw_inheritance(ax, 12.2, 0.3, 12.2, -0.3)

    # Legend
    legend_y = 9.2
    ax.text(0.3, legend_y, "Legend:", fontsize=8, fontweight="bold", color="#01411C")
    ax.plot([1.2, 1.8], [legend_y - 0.05, legend_y - 0.05], color="#046A38", lw=1.5)
    tri = Polygon([(1.95, legend_y - 0.05), (1.77, legend_y - 0.22), (2.13, legend_y - 0.22)],
                  closed=True, fill=False, edgecolor="#046A38", lw=1.2)
    ax.add_patch(tri)
    ax.text(2.2, legend_y - 0.08, "Inheritance (extends)", fontsize=7)
    diamond = Polygon([(4.0, legend_y), (4.12, legend_y - 0.08), (4.0, legend_y - 0.16), (3.88, legend_y - 0.08)],
                      closed=True, facecolor="#01411C", edgecolor="#01411C")
    ax.add_patch(diamond)
    ax.plot([4.15, 4.6], [legend_y - 0.08, legend_y - 0.08], color="#01411C", lw=1.2)
    ax.text(4.7, legend_y - 0.08, "Composition (owns)", fontsize=7)
    ax.annotate("", xy=(7.0, legend_y - 0.08), xytext=(6.3, legend_y - 0.08),
                arrowprops=dict(arrowstyle="-|>", color="#666", lw=1, linestyle="dashed"))
    ax.text(7.1, legend_y - 0.08, "Dependency (uses)", fontsize=7)
    ax.annotate("", xy=(9.3, legend_y - 0.08), xytext=(8.6, legend_y - 0.08),
                arrowprops=dict(arrowstyle="-|>", color="#333", lw=1))
    ax.text(9.4, legend_y - 0.08, "Association", fontsize=7)

    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def create_architecture_diagram(path: Path):
    fig, ax = plt.subplots(figsize=(10, 8))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 10)
    ax.axis("off")
    ax.set_title("PakBzer — Layered System Architecture", fontsize=13, fontweight="bold", color="#01411C", pad=12)

    layers = [
        (1.0, 8.5, 8.0, 0.9, "Presentation Layer", "Controllers + Thymeleaf HTML\nHome, Product, Cart, Auth, Checkout, Review"),
        (1.0, 7.0, 8.0, 0.9, "Business Logic Layer", "Services: ProductService, UserService, ReviewService\nCart (session), StripeService"),
        (1.0, 5.5, 8.0, 0.9, "Data Access Layer", "Repositories: Product, User, Review\n(Spring Data JPA interfaces)"),
        (1.0, 4.0, 8.0, 0.9, "Domain Model Layer", "Entities: Product, User, Review, CartItem\nBaseEntity, Category, Role enums"),
        (1.0, 2.5, 8.0, 0.9, "Persistence", "H2 Embedded Database (./data/pakbzer)"),
    ]
    for x, y, w, h, title, body in layers:
        box = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02", linewidth=1.5,
                             edgecolor="#046A38", facecolor="#f1f6f2")
        ax.add_patch(box)
        ax.text(x + w / 2, y + h - 0.22, title, ha="center", fontsize=9, fontweight="bold", color="#01411C")
        ax.text(x + w / 2, y + 0.15, body, ha="center", va="bottom", fontsize=7)

    # External systems
    for x, label in [(0.3, "Browser\n(HTTP)"), (8.8, "Google\nOAuth2"), (8.8, "Stripe\nPayments")]:
        pass
    ax.text(0.5, 9.5, "Browser\n(HTTP)", ha="center", fontsize=7, bbox=dict(boxstyle="round", facecolor="#fff8e1", edgecolor="#c9a227"))
    ax.text(9.5, 8.9, "Google OAuth2", ha="center", fontsize=7, bbox=dict(boxstyle="round", facecolor="#fff8e1", edgecolor="#c9a227"))
    ax.text(9.5, 7.4, "Stripe API", ha="center", fontsize=7, bbox=dict(boxstyle="round", facecolor="#fff8e1", edgecolor="#c9a227"))

    # Vertical flow arrows between layers
    for y_from, y_to in [(8.5, 7.9), (7.0, 6.4), (5.5, 4.9), (4.0, 3.4)]:
        ax.annotate("", xy=(5, y_to), xytext=(5, y_from),
                    arrowprops=dict(arrowstyle="-|>", color="#046A38", lw=2))
    ax.text(5.15, 8.2, "HTTP request", fontsize=6.5, color="#046A38")
    ax.text(5.15, 6.7, "service call", fontsize=6.5, color="#046A38")
    ax.text(5.15, 5.2, "repository call", fontsize=6.5, color="#046A38")
    ax.text(5.15, 3.7, "JPA / Hibernate", fontsize=6.5, color="#046A38")

    # Cross-cutting security
    sec = FancyBboxPatch((0.2, 0.4), 9.6, 1.6, boxstyle="round,pad=0.02", linewidth=1.2,
                         edgecolor="#01411C", facecolor="#e8f4ec", linestyle="--")
    ax.add_patch(sec)
    ax.text(5, 1.55, "Cross-Cutting: Security Layer", ha="center", fontsize=9, fontweight="bold", color="#01411C")
    ax.text(5, 0.85, "SecurityConfig | CustomUserDetailsService | CustomOidcUserService\nConfig: PasswordConfig, StripeProperties, DataSeeder",
            ha="center", fontsize=7)

    # External deps
    draw_dependency(ax, 9.2, 8.9, 8.5, 8.2)
    draw_dependency(ax, 9.2, 7.4, 8.5, 7.2)

    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def create_sequence_diagram(path: Path):
    fig, ax = plt.subplots(figsize=(11, 6))
    ax.set_xlim(0, 11)
    ax.set_ylim(0, 6)
    ax.axis("off")
    ax.set_title('PakBzer — Sequence: "Add to Cart"', fontsize=12, fontweight="bold", color="#01411C", pad=10)

    actors = ["Browser", "CartController", "ProductService", "ProductRepository", "Cart"]
    xs = [1, 3, 5, 7, 9]
    for x, name in zip(xs, actors):
        ax.text(x, 5.5, name, ha="center", fontsize=8, fontweight="bold",
                bbox=dict(boxstyle="round", facecolor="#f1f6f2", edgecolor="#046A38"))
        ax.plot([x, x], [0.5, 5.2], color="#ccc", linestyle="--", linewidth=1)

    steps = [
        (0, 1, 5.0, 4.7, "POST /cart/add"),
        (1, 2, 4.5, 4.2, "findById(id)"),
        (2, 3, 4.0, 3.7, "findById(id)"),
        (3, 2, 3.5, 3.2, "Product"),
        (2, 1, 3.0, 2.7, "Product"),
        (1, 4, 2.5, 2.2, "add(product, qty)"),
        (4, 1, 2.0, 1.7, "ok"),
        (1, 0, 1.5, 1.2, "redirect /cart"),
    ]
    for i_from, i_to, y1, y2, label in steps:
        style = "-|>"
        ls = "dashed" if i_from > i_to and i_from - i_to > 1 else "solid"
        ax.annotate(
            "", xy=(xs[i_to], y2), xytext=(xs[i_from], y1),
            arrowprops=dict(arrowstyle=style, color="#046A38", lw=1.2, linestyle=ls),
        )
        ax.text((xs[i_from] + xs[i_to]) / 2, (y1 + y2) / 2 + 0.08, label, fontsize=6.5, ha="center")

    fig.tight_layout()
    fig.savefig(path, dpi=180, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def img_to_base64(path: Path) -> str:
    return base64.b64encode(path.read_bytes()).decode("ascii")


def update_html_doc(class_png, arch_png, seq_png):
    try:
        html = DOC_HTML.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        print(f"Skipping legacy .doc text refresh: {DOC_HTML} is not UTF-8 text.")
        return
    except FileNotFoundError:
        print(f"Skipping legacy .doc text refresh: {DOC_HTML} not found.")
        return
    diagram_section = f"""
<h3>2.1 UML Class Diagram</h3>
<p>The diagram below shows inheritance from <code>BaseEntity</code>, composition between
<code>Product</code> and <code>Review</code> / <code>Cart</code> and <code>CartItem</code>,
associations to enums, interface realization (<code>ProductServiceImpl</code>), and dependencies.</p>
<p style="text-align:center;"><img src="data:image/png;base64,{img_to_base64(class_png)}" width="680" alt="UML Class Diagram"/></p>

<h3>2.2 Layered Architecture Diagram</h3>
<p>Request flow moves downward through presentation, business, data-access, and domain layers.
Security and configuration cut across all layers. External integrations: Google OAuth2 and Stripe.</p>
<p style="text-align:center;"><img src="data:image/png;base64,{img_to_base64(arch_png)}" width="580" alt="System Architecture Diagram"/></p>

<h3>2.3 Add-to-Cart Sequence Diagram</h3>
<p>End-to-end flow when a user adds a product to the shopping cart, showing layer interactions
with arrows indicating call direction.</p>
<p style="text-align:center;"><img src="data:image/png;base64,{img_to_base64(seq_png)}" width="620" alt="Add to Cart Sequence Diagram"/></p>
"""
    if "<h3>2.1 UML Class Diagram</h3>" in html:
        html = re.sub(
            r"<h3>2\.1 UML Class Diagram</h3>.*?(?=<h2>3\. The Four Pillars)",
            diagram_section,
            html,
            flags=re.DOTALL,
        )
    else:
        html = html.replace("<h2>3. The Four Pillars", diagram_section + "\n<h2>3. The Four Pillars")

    DOC_HTML.write_text(html, encoding="utf-8")


def create_docx(class_png, arch_png, seq_png):
    """Build a proper .docx from REPORT content with embedded diagram images."""
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("PakBzer", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("A Pakistani E-Commerce Web Application")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2 = doc.add_paragraph("Object-Oriented Programming Project Report")
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub2.runs[0].bold = True
    doc.add_paragraph("Language: Java 17 | Framework: Spring Boot 3 | Paradigm: OOP").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Abstract", level=1)
    doc.add_paragraph(
        "PakBzer is a full-stack online clothing store for the Pakistani market. It demonstrates "
        "Encapsulation, Inheritance, Abstraction, and Polymorphism in a layered Spring Boot application."
    )

    doc.add_heading("2. System Architecture", level=1)
    doc.add_paragraph(
        "PakBzer follows a classic layered (n-tier) architecture. Each layer has a single responsibility "
        "and communicates with the layer directly beneath it."
    )

    doc.add_heading("2.1 UML Class Diagram", level=2)
    doc.add_paragraph(
        "Inheritance (hollow triangle): Product, User, Review extend BaseEntity. "
        "Composition (filled diamond): Product owns Reviews; Cart owns CartItems. "
        "Association: Product→Category, User→Role. "
        "Dependency (dashed arrow): ProductServiceImpl uses ProductRepository. "
        "Realization (dashed triangle): ProductServiceImpl implements ProductService."
    )
    doc.add_picture(str(class_png), width=Inches(6.5))

    doc.add_heading("2.2 Layered Architecture Diagram", level=2)
    doc.add_paragraph(
        "HTTP requests flow Browser → Controllers → Services → Repositories → H2 Database. "
        "Security (OAuth2, form login) and configuration cross-cut all layers."
    )
    doc.add_picture(str(arch_png), width=Inches(5.5))

    doc.add_heading("2.3 Add-to-Cart Sequence Diagram", level=2)
    doc.add_paragraph(
        "Shows the presentation → business → data-access call chain when adding a product to the session cart."
    )
    doc.add_picture(str(seq_png), width=Inches(6.0))

    doc.add_heading("3. OOP Pillars & Relationships", level=1)
    pillars = [
        ("Encapsulation", "Private fields with getters/setters; business rules in methods (getAverageRating, getLineTotal)."),
        ("Inheritance", "Product, User, Review extend abstract BaseEntity for shared id and timestamps."),
        ("Abstraction", "ProductService interface hides implementation; BaseEntity is abstract."),
        ("Polymorphism", "ProductServiceImpl implements ProductService; CustomOidcUserService overrides OidcUserService."),
        ("Composition", "Cart composes CartItems; Product composes Reviews (cascade ALL, orphanRemoval)."),
        ("Association", "Product linked to Category enum; User linked to Role enum; Review linked to Product."),
    ]
    for name, desc in pillars:
        p = doc.add_paragraph()
        p.add_run(name + ": ").bold = True
        p.add_run(desc)

    doc.add_heading("4. Conclusion", level=1)
    doc.add_paragraph(
        "PakBzer applies OOP principles in a modular Spring Boot e-commerce system that is "
        "reusable, testable, and easy to extend."
    )

    doc.save(DOCX_OUT)


def main():
    DIAGRAMS.mkdir(exist_ok=True)
    class_png = DIAGRAMS / "uml_class_diagram.png"
    arch_png = DIAGRAMS / "system_architecture.png"
    seq_png = DIAGRAMS / "add_to_cart_sequence.png"

    print("Generating diagrams...")
    create_class_diagram(class_png)
    create_architecture_diagram(arch_png)
    create_sequence_diagram(seq_png)

    print("Updating HTML .doc report...")
    update_html_doc(class_png, arch_png, seq_png)

    print("Creating .docx report...")
    create_docx(class_png, arch_png, seq_png)

    print(f"Done.\n  Diagrams: {DIAGRAMS}\n  Updated: {DOC_HTML}\n  Created: {DOCX_OUT}")


if __name__ == "__main__":
    main()
