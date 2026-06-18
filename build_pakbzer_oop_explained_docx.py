from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "PakBzer_OOP_Detailed_Explanation.docx"
UMl_DIAGRAM = ROOT / "diagrams" / "uml_class_diagram.png"
ARCH_DIAGRAM = ROOT / "diagrams" / "system_architecture.png"
UMl_DIAGRAM_PRINT = ROOT / "diagrams" / "uml_class_diagram_print.jpg"
ARCH_DIAGRAM_PRINT = ROOT / "diagrams" / "system_architecture_print.jpg"

SKILL_SCRIPTS = Path(
    r"C:\Users\user\.codex\plugins\cache\openai-primary-runtime\documents\26.614.11602\skills\documents\scripts"
)
if str(SKILL_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(SKILL_SCRIPTS))

from table_geometry import apply_table_geometry  # noqa: E402


GREEN = RGBColor(0x01, 0x41, 0x1C)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
TEXT = RGBColor(0x00, 0x00, 0x00)
MUTED = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "F2F4F7"
NOTE_FILL = "EEF5EE"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, *, bold: bool = False, color=TEXT, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = color


def format_paragraph(paragraph, *, before=0, after=6, line=1.1, align=WD_ALIGN_PARAGRAPH.LEFT):
    paragraph.alignment = align
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def set_run_font(run, *, name="Calibri", size=11, bold=False, italic=False, color=TEXT):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def add_paragraph(doc: Document, text: str = "", *, style=None, before=0, after=6, line=1.1, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        set_run_font(run)
    format_paragraph(p, before=before, after=after, line=line, align=align)
    return p


def add_bold_label_paragraph(doc: Document, label: str, text: str, *, after=6):
    p = doc.add_paragraph()
    format_paragraph(p, after=after)
    r1 = p.add_run(label)
    set_run_font(r1, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    set_run_font(run, size=9.5, italic=True, color=GREEN)
    return p


def add_centered_image(doc: Document, image_path: Path, width_inches: float):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.add_run().add_picture(str(image_path), width=Inches(width_inches))
    return p


def flatten_png_for_print(source: Path, target: Path, max_dim: int = 1600) -> Path:
    """Flatten transparent images onto white so LibreOffice can export them reliably."""

    if target.exists() and target.stat().st_mtime >= source.stat().st_mtime:
        return target

    img = Image.open(source)
    img = img.copy()
    if img.mode != "RGB":
        base = Image.new("RGB", img.size, (255, 255, 255))
        if img.mode in ("RGBA", "LA"):
            base.paste(img, mask=img.getchannel("A"))
        else:
            base.paste(img.convert("RGBA"))
        img = base

    img.thumbnail((max_dim, max_dim), Image.Resampling.LANCZOS)
    img.save(target, format="JPEG", quality=95, optimize=True)
    return target


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int], title: str | None = None):
    if title:
        add_heading(doc, title, level=2)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=GREEN, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[idx], HEADER_FILL)
        table.rows[0].cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=10.5)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.TOP

    apply_table_geometry(
        table,
        widths,
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    return table


def set_document_theme(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1.font.size = Pt(16)
    h1.font.bold = False
    h1.font.color.rgb = BLUE
    h1.paragraph_format.space_before = Pt(16)
    h1.paragraph_format.space_after = Pt(8)
    h1.paragraph_format.line_spacing = 1.1

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2.font.size = Pt(13)
    h2.font.bold = False
    h2.font.color.rgb = BLUE
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.line_spacing = 1.1

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Calibri"
    h3._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h3._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h3.font.size = Pt(12)
    h3.font.bold = False
    h3.font.color.rgb = DARK_BLUE
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.line_spacing = 1.1


def cover_page(doc: Document):
    for _ in range(4):
        add_paragraph(doc, "", after=0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("PakBzer OOP Project Explanation")
    set_run_font(run, size=24, bold=True, color=GREEN)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("A detailed guide to the four OOP pillars, UML arrows, and layered architecture used in the project")
    set_run_font(run, size=12.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("Java 17 | Spring Boot 3 | Thymeleaf | H2 | Stripe | Google OAuth2")
    set_run_font(run, size=11.5, bold=True, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(f"Prepared on {date.today().strftime('%d %B %Y')}")
    set_run_font(run, size=10.5, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("This document explains how PakBzer uses object-oriented design in a real e-commerce system.")
    set_run_font(run, size=11, italic=True, color=TEXT)

    doc.add_page_break()


def add_project_overview(doc: Document):
    add_heading(doc, "1. Project Overview", level=1)
    add_paragraph(
        doc,
        "PakBzer is a Pakistani fashion e-commerce web application built with Java 17 and Spring Boot 3. "
        "It lets users browse product categories, search items, add products to a cart, register or sign in "
        "with email or Google OAuth2, write reviews, and complete checkout through Stripe. "
        "Because the project is structured around domain objects, services, repositories, and controllers, "
        "it is a strong example of object-oriented programming in a practical web system.",
    )
    add_paragraph(
        doc,
        "The project is also a good teaching example because the code clearly separates responsibilities. "
        "The model classes hold the core data, the services hold business rules, the repositories hide database "
        "details, and the controllers handle web requests. That separation makes it easier to explain where each "
        "OOP concept appears in the code.",
    )

    add_table(
        doc,
        ["OOP pillar", "Where it appears in PakBzer"],
        [
            ["Encapsulation", "Product, User, Review, Cart, and CartItem keep their fields private and expose controlled behavior through methods."],
            ["Inheritance", "BaseEntity provides shared id and timestamps to Product, User, and Review; CustomOidcUserService extends OidcUserService."],
            ["Abstraction", "ProductService is an interface, BaseEntity is abstract, and Spring Data repositories hide SQL details."],
            ["Polymorphism", "ProductServiceImpl implements ProductService, while CustomOidcUserService overrides loadUser() at runtime."],
        ],
        [2100, 7260],
        title="2. Quick map of the four OOP pillars",
    )


def add_encapsulation_section(doc: Document):
    add_heading(doc, "3. Encapsulation", level=1)
    add_paragraph(
        doc,
        "Encapsulation means an object protects its own state and offers methods that control how that state changes. "
        "In PakBzer, this appears in both the domain model and the service layer. "
        "Instead of letting controllers calculate totals or ratings on their own, the objects that own the data "
        "perform those calculations internally.",
    )
    add_paragraph(
        doc,
        "For example, Product calculates its own average rating and in-stock status, Cart calculates item totals and the grand total, "
        "and CartItem computes its own line total. UserService also encapsulates registration rules by rejecting duplicate emails "
        "and hashing passwords before saving the account. "
        "That keeps the rules close to the data they protect.",
    )

    add_paragraph(doc, "Key PakBzer examples:", after=4)
    add_bold_label_paragraph(doc, "Product: ", "getAverageRating(), getReviewCount(), and isInStock() are calculated inside the Product object.")
    add_bold_label_paragraph(doc, "Cart: ", "add(), updateQuantity(), remove(), and getTotalPrice() control all cart changes.")
    add_bold_label_paragraph(doc, "CartItem: ", "getLineTotal() computes unit price times quantity instead of storing a duplicate value.")
    add_bold_label_paragraph(doc, "UserService: ", "register() checks for duplicate emails and stores BCrypt-hashed passwords.")


def add_inheritance_section(doc: Document):
    add_heading(doc, "4. Inheritance", level=1)
    add_paragraph(
        doc,
        "Inheritance lets one class reuse the fields and behavior of another class. "
        "In PakBzer, the abstract BaseEntity class is the inheritance root for all persistent entities. "
        "It defines the shared id, createdAt, and updatedAt fields and uses JPA lifecycle hooks to set timestamps automatically.",
    )
    add_paragraph(
        doc,
        "Product, User, and Review extend BaseEntity, so each one gets the same identity and audit behavior without copy-pasting code. "
        "The project also uses inheritance outside the model layer: CustomOidcUserService extends Spring Security's OidcUserService "
        "so the Google login flow can be customized for PakBzer's own user database.",
    )

    add_bold_label_paragraph(doc, "BaseEntity: ", "shared id and timestamp logic for every persistent object.")
    add_bold_label_paragraph(doc, "Product / User / Review: ", "domain classes that inherit the common audit fields.")
    add_bold_label_paragraph(doc, "CustomOidcUserService: ", "specializes the Google sign-in flow by extending OidcUserService.")


def add_abstraction_polymorphism_section(doc: Document):
    add_heading(doc, "5. Abstraction and Polymorphism", level=1)
    add_paragraph(
        doc,
        "Abstraction hides implementation details behind a contract. In PakBzer, ProductService is an interface that tells the rest of "
        "the system what catalogue operations exist without exposing how those operations are implemented. "
        "BaseEntity is also abstract: it exists only as a shared parent and cannot be instantiated directly.",
    )
    add_paragraph(
        doc,
        "Polymorphism is what lets the application use those abstractions flexibly. Controllers can depend on ProductService, but at runtime "
        "the actual object is ProductServiceImpl. The same idea appears in CustomOidcUserService, where overriding loadUser() changes the "
        "behavior of the Google login flow without changing Spring Security's external contract.",
    )

    add_paragraph(doc, "Practical effect:", after=4)
    add_bold_label_paragraph(doc, "ProductService interface: ", "callers see findAll(), findById(), findByCategory(), search(), and save().")
    add_bold_label_paragraph(doc, "ProductServiceImpl: ", "provides the real behavior by delegating to ProductRepository.")
    add_bold_label_paragraph(doc, "CustomUserDetailsService: ", "implements Spring Security's UserDetailsService for local email/password login.")
    add_bold_label_paragraph(doc, "CustomOidcUserService: ", "overrides loadUser() so Google-authenticated users are also provisioned inside PakBzer.")


def add_relationships_section(doc: Document):
    add_heading(doc, "6. UML Relationships and Arrow Meaning", level=1)
    add_paragraph(
        doc,
        "The UML diagram in this project uses standard class-diagram symbols. "
        "If you want the short answer to 'is this dependent or independent?', the rule is simple: "
        "composition is dependent, association is independent, and dependency is a temporary 'uses' relationship.",
    )
    add_paragraph(
        doc,
        "The diagram also distinguishes 'is-a' relationships from 'uses' relationships. "
        "A hollow triangle at the parent side means inheritance or interface realization. "
        "A filled diamond means strong ownership. "
        "A dashed arrow means one class depends on another for a specific operation.",
    )

    add_table(
        doc,
        ["Relationship", "Symbol in the diagram", "Dependent?", "PakBzer example"],
        [
            ["Inheritance / generalization", "Solid line with hollow triangle", "Type relationship, not ownership", "Product, User, and Review extend BaseEntity."],
            ["Realization / implements", "Dashed line with hollow triangle", "No ownership", "ProductServiceImpl implements ProductService."],
            ["Composition", "Filled diamond", "Yes, strong dependency", "Cart owns CartItem objects; Product owns Review objects."],
            ["Association", "Plain line", "Mostly independent", "Product is associated with Category; User is associated with Role."],
            ["Dependency", "Dashed arrow", "Weak / temporary use", "ProductServiceImpl uses ProductRepository; StripeService uses Stripe SDK."],
        ],
        [1700, 2500, 1700, 3460],
        title="7. How to read each arrow",
    )

    add_heading(doc, "7.1 What each arrow means in simple words", level=2)
    add_bold_label_paragraph(
        doc,
        "Composition: ",
        "the child object belongs to the parent object. If the parent goes away, the child is not meant to stand alone in the model.",
    )
    add_bold_label_paragraph(
        doc,
        "Association: ",
        "the objects are linked, but neither side owns the other's lifecycle. The link is important, but the classes can exist independently.",
    )
    add_bold_label_paragraph(
        doc,
        "Dependency: ",
        "one class needs another class to perform a task. The relationship is temporary and weaker than composition or association.",
    )
    add_bold_label_paragraph(
        doc,
        "Inheritance / realization: ",
        "the child class specializes a parent class or fulfills an interface contract. This is the 'is-a' family of relationships.",
    )

    add_heading(doc, "7.2 PakBzer relationships in plain English", level=2)
    add_paragraph(
        doc,
        "Product -> Review is a composition relationship because a product owns its reviews in the data model. "
        "Cart -> CartItem is also composition because the cart stores and manages its own line items. "
        "Product -> Category and User -> Role are association-style links because they are modeled as enum references; "
        "the category or role values exist independently of any one product or user. "
        "ProductServiceImpl -> ProductRepository is a dependency because the service calls repository methods when it needs data, "
        "but it does not own the repository. "
        "ProductServiceImpl -> ProductService is realization because the class fulfills the interface contract.",
    )

    add_paragraph(
        doc,
        "The CustomOidcUserService -> OidcUserService arrow is inheritance. "
        "That means PakBzer does not merely 'use' the parent class; it specializes it. "
        "The override of loadUser() lets the project keep Spring Security's behavior while adding the extra step of creating "
        "a PakBzer user record after Google login.",
    )


def add_uml_figure(doc: Document):
    add_heading(doc, "8. UML Class Diagram", level=1)
    add_paragraph(
        doc,
        "The following diagram shows the core PakBzer domain classes and the most important service and repository relationships. "
        "It is the best visual reference for understanding the arrows described above.",
    )
    uml_path = flatten_png_for_print(UMl_DIAGRAM, UMl_DIAGRAM_PRINT)
    if uml_path.exists():
        add_centered_image(doc, uml_path, 6.5)
        add_caption(doc, "Figure 1. PakBzer UML class diagram showing inheritance, composition, association, realization, and dependency.")
    else:
        add_paragraph(doc, "UML diagram image not found.", after=8)


def add_architecture_section(doc: Document):
    add_heading(doc, "9. Layered Architecture", level=1)
    add_paragraph(
        doc,
        "PakBzer is built with a layered architecture, which keeps the user interface, business rules, data access, and persistence concerns separate. "
        "That structure makes the code easier to test, easier to explain, and easier to extend.",
    )
    add_table(
        doc,
        ["Layer", "Main classes", "Role in PakBzer"],
        [
            ["Presentation", "HomeController, ProductController, CartController, AuthController, ReviewController, CheckoutController", "Handles browser requests and returns Thymeleaf views."],
            ["Business logic", "ProductService, ProductServiceImpl, UserService, ReviewService, Cart, StripeService", "Contains the rules for search, registration, cart handling, reviews, and checkout."],
            ["Data access", "ProductRepository, UserRepository, ReviewRepository", "Wraps database queries through Spring Data JPA."],
            ["Domain model", "BaseEntity, Product, User, Review, CartItem, Category, Role", "Stores the core business data and object behavior."],
            ["Persistence and security", "H2 database, SecurityConfig, CustomUserDetailsService, CustomOidcUserService, PasswordConfig, StripeProperties", "Stores data and protects the application with login, OAuth2, and payment configuration."],
        ],
        [1500, 2900, 4960],
        title="10. Main layers in the system architecture",
    )
    arch_path = flatten_png_for_print(ARCH_DIAGRAM, ARCH_DIAGRAM_PRINT)
    if arch_path.exists():
        add_centered_image(doc, arch_path, 6.5)
        add_caption(doc, "Figure 2. PakBzer layered system architecture from the browser down to the H2 database.")
    else:
        add_paragraph(doc, "Architecture diagram image not found.", after=8)

    add_paragraph(
        doc,
        "The arrow flow is top-down: browser request -> controller -> service -> repository -> database. "
        "The security layer and configuration classes sit across the stack, because authentication, password hashing, Google OAuth2, and Stripe "
        "do not belong to only one application layer. "
        "That cross-cutting design is another example of good separation of concerns.",
    )


def add_class_map_section(doc: Document):
    add_heading(doc, "11. Main Classes and Their OOP Role", level=1)
    add_paragraph(
        doc,
        "This table ties the codebase back to the OOP explanation so the diagram is not just visual decoration. "
        "It shows where the most important classes live and what design idea each one demonstrates.",
    )
    add_table(
        doc,
        ["Package", "Key classes", "Why they matter"],
        [
            ["model", "BaseEntity, Product, User, Review, CartItem, Category, Role", "Core object model. This is where encapsulation, inheritance, composition, and enums are most visible."],
            ["service", "ProductService, ProductServiceImpl, UserService, ReviewService, Cart, StripeService", "Business logic and object behavior. This is where abstraction and polymorphism are strongest."],
            ["repository", "ProductRepository, UserRepository, ReviewRepository", "Database abstraction through Spring Data JPA interfaces."],
            ["security", "SecurityConfig, CustomUserDetailsService, CustomOidcUserService", "Authentication and authorization, including Google sign-in customization."],
            ["controller", "HomeController, ProductController, CartController, AuthController, ReviewController, CheckoutController, GlobalControllerAdvice", "Web entry points that connect HTTP requests to the application logic."],
        ],
        [1800, 3100, 4460],
        title="11. Package-level class map",
    )


def add_end_to_end_example(doc: Document):
    add_heading(doc, "12. End-to-End Example: Adding a Product to the Cart", level=1)
    add_paragraph(
        doc,
        "A simple shopping action shows several OOP ideas working together. When the user clicks Add to Cart, the controller receives the request, "
        "the service layer loads the Product through ProductRepository, and the Cart object updates its own list of CartItem objects. "
        "The controller does not calculate totals or decide how items are merged; the Cart object owns those rules.",
    )
    add_paragraph(
        doc,
        "This single user action demonstrates encapsulation because the cart hides its internal list; abstraction because the controller talks to the service "
        "contract rather than SQL; and polymorphism because Spring can inject the correct service implementation behind the interface. "
        "It is a good real-world example of why OOP matters in PakBzer.",
    )


def add_conclusion(doc: Document):
    add_heading(doc, "13. Conclusion", level=1)
    add_paragraph(
        doc,
        "PakBzer uses OOP in a practical and understandable way. Encapsulation keeps state safe, inheritance removes duplicate audit fields, "
        "abstraction hides implementation details, and polymorphism makes the system flexible. "
        "The UML diagram shows how those ideas connect through inheritance, composition, association, dependency, and interface realization.",
    )
    add_paragraph(
        doc,
        "If you want a one-line summary: the project is not just an online store, it is a clean example of how object-oriented design can organize a real Spring Boot application.",
    )


def build_document():
    doc = Document()
    set_document_theme(doc)
    cover_page(doc)

    add_project_overview(doc)
    add_encapsulation_section(doc)
    add_inheritance_section(doc)
    add_abstraction_polymorphism_section(doc)

    doc.add_page_break()
    add_relationships_section(doc)

    doc.add_page_break()
    add_uml_figure(doc)
    doc.add_page_break()
    add_architecture_section(doc)

    doc.add_page_break()
    add_class_map_section(doc)
    add_end_to_end_example(doc)
    add_conclusion(doc)

    doc.save(OUT)


if __name__ == "__main__":
    build_document()
    print(f"Created {OUT}")
